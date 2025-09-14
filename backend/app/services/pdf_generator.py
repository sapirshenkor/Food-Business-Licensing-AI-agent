"""
PDF and DOCX generation service for reports
"""

import markdown
from weasyprint import HTML, CSS
from docx import Document
from docx.shared import Inches
import os
import tempfile
from typing import Optional

class DocumentGenerator:
    def __init__(self):
        self.temp_dir = tempfile.gettempdir()
    
    def markdown_to_pdf(self, markdown_text: str, filename: str = None) -> str:
        """Convert markdown report to PDF"""
        try:
            # Convert markdown to HTML
            html_content = markdown.markdown(markdown_text, extensions=['tables', 'nl2br'])
            
            # Add Hebrew RTL CSS styling
            html_with_style = f"""
            <!DOCTYPE html>
            <html dir="rtl" lang="he">
            <head>
                <meta charset="UTF-8">
                <style>
                    body {{
                        font-family: 'Arial', 'Helvetica', sans-serif;
                        direction: rtl;
                        text-align: right;
                        line-height: 1.6;
                        margin: 20px;
                        color: #333;
                    }}
                    h1, h2, h3 {{
                        color: #2c3e50;
                        border-bottom: 2px solid #3498db;
                        padding-bottom: 10px;
                    }}
                    h1 {{ font-size: 28px; }}
                    h2 {{ font-size: 22px; }}
                    h3 {{ font-size: 18px; }}
                    strong {{ color: #e74c3c; }}
                    ul, ol {{ padding-right: 20px; }}
                    li {{ margin-bottom: 5px; }}
                    .report-header {{
                        background: #3498db;
                        color: white;
                        padding: 15px;
                        border-radius: 5px;
                        margin-bottom: 20px;
                        text-align: center;
                    }}
                </style>
            </head>
            <body>
                <div class="report-header">
                    <h1>דוח רישוי עסקים</h1>
                    <p>נוצר ב-{datetime.now().strftime('%d/%m/%Y %H:%M')}</p>
                </div>
                {html_content}
            </body>
            </html>
            """
            
            # Generate filename if not provided
            if not filename:
                from datetime import datetime
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"business_report_{timestamp}.pdf"
            
            # Generate PDF
            pdf_path = os.path.join(self.temp_dir, filename)
            HTML(string=html_with_style).write_pdf(pdf_path)
            
            return pdf_path
            
        except Exception as e:
            print(f"❌ PDF generation error: {e}")
            return None
    
    def markdown_to_docx(self, markdown_text: str, filename: str = None) -> str:
        """Convert markdown report to DOCX"""
        try:
            # Create new document
            doc = Document()
            
            # title
            title = doc.add_heading('דוח רישוי עסקים', 0)
            title.alignment = 2  # Right align for Hebrew
            
            # Add creation date
            from datetime import datetime
            date_para = doc.add_paragraph(f'נוצר ב-{datetime.now().strftime("%d/%m/%Y %H:%M")}')
            date_para.alignment = 2
            
            # Process markdown content line by line
            lines = markdown_text.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Headers
                if line.startswith('# '):
                    heading = doc.add_heading(line[2:], 1)
                    heading.alignment = 2
                elif line.startswith('## '):
                    heading = doc.add_heading(line[3:], 2)
                    heading.alignment = 2
                elif line.startswith('### '):
                    heading = doc.add_heading(line[4:], 3)
                    heading.alignment = 2
                # Lists
                elif line.startswith('- '):
                    para = doc.add_paragraph(line[2:], style='List Bullet')
                    para.alignment = 2
                # Bold text
                elif line.startswith('**') and line.endswith('**'):
                    para = doc.add_paragraph()
                    para.alignment = 2
                    run = para.add_run(line[2:-2])
                    run.bold = True
                # Regular text
                else:
                    para = doc.add_paragraph(line)
                    para.alignment = 2
            
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"business_report_{timestamp}.docx"
            
            # Save document
            docx_path = os.path.join(self.temp_dir, filename)
            doc.save(docx_path)
            
            return docx_path
            
        except Exception as e:
            print(f"❌ DOCX generation error: {e}")
            return None
    
    def clean_temp_files(self, file_path: str):
        """Clean up temporary files"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            print(f"⚠️ Error cleaning temp file: {e}")